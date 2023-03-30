#coding: utf8

import re

from docx import Document
from docx.shared import Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Text

## Textual Content

def add_content(line, style = "Normal"):
    words = line.split(" ")

    p = document.add_paragraph(style = style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    #print(words)

    for word in words:
        word += " "

        if (word[0:2] == "**"):
            word = word.replace("**", "")

            p.add_run(word).bold = True
        else:
            p.add_run(word)

## List Bullet

def add_list_bullet(line):
    line = line.replace("- ", "")

    add_content(line, style = "List Bullet")

# Image

def add_image(line):
    filename = re.sub("(\!|\[|\]|\(|\))", "", line)

    document.add_picture(filename, width = Cm(15))

# Titles

## Heading Title

def add_heading_title(line):
    level = 1

    if line[0:2] == "##": level = 2

    line = re.sub("#\s?", "", line)

    document.add_heading(line, level)

    document.add_paragraph()

## Bold Title

def add_bold_title(line):
    line = line.replace("*", "")

    p = document.add_paragraph()
    p.add_run(line).bold = True

# Table

## Define Table

def add_table(line):
    num_cols = line.count("|") - 1

    table = document.add_table(rows = 0, cols = num_cols)
    table.style = "Table Grid"

    document.add_paragraph()

    return table

## Add Row

def add_row(table, line):
    if line[0] == "|":
        line = re.sub("(^\|)|(\|$)", "", line)

        cols = line.split("|")

        row_cells = table.add_row().cells

        num_cols = len(cols)

        for index in range(0, num_cols):
            col = cols[index]
            row = row_cells[index]

            row.text = col
    else:
        document.add_paragraph()

# Code

def add_code(line):
    words = line.split(" ")

    p = document.add_paragraph(style = "No Spacing")

    for word in words:
        word += " "

        run = p.add_run(word)

        font = run.font

        if re.match("^[A-Z]", word):    blue = 255
        else:   blue = 0

        font.color.rgb = RGBColor(0, 0, blue)

# Doc Parts

def add_part(filename):
    table = None
    is_code = False

    with open(f"docs/parts/{filename}", "r") as file:
        for line in file:
            line = line.strip()

            if len(line) > 0:
                first = line[0]

                if first == "#":
                    add_heading_title(line)
                elif first == "*":
                    add_bold_title(line)
                elif first == "-":
                    add_list_bullet(line)
                elif first == "!":
                    add_image(line)
                elif first == "|":
                    if re.match("^\|{3,5}$", line):
                        table = add_table(line)
                    else:
                        add_row(table, line)
                elif first == "`":
                    if re.match("`{3}sql", line):   is_code = True
                    else:   is_code = False

                    document.add_paragraph()
                elif is_code:
                    add_code(line)
                else:
                    add_content(line)

        file.close()

# Convert MD Files to Docx

def convert_md_files():
    with open("docs/parts.txt", "r") as file:
        for line in file:
            line = line.strip()

            if line[0] != "#":
                filename = line + ".md"

                add_part(filename)

        file.close()

    document.save("result/compilation.docx")

    print("Compilation du dossier effectuée")

# Docx File

document = Document()

# Set Font & Style

style = document.styles["Normal"]

font = style.font
font.name = "Calibri"

# Compilation

convert_md_files()