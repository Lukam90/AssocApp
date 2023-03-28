#coding: utf8

import re

from docx import Document
from docx.shared import Cm

# Doc Table

def add_table(filename, num_cols):
    global document

    table = document.add_table(rows = 1, cols = num_cols)
    table.style = "Table Grid"

    with open(f"docs/tables/{filename}", "r") as file:
        for line in file:
            line = line.strip()

            cols = line.split(";")

            row_cells = table.add_row().cells

            for index in range(0, num_cols):
                col = cols[index]
                row = row_cells[index]

                row.text = col

    file.close()

# Doc Image

def add_image(filename):
    global document

    document.add_picture(f"docs/{filename}", width = Cm(15))

# Doc Textual Content

def add_content(line, style = "Normal"):
    global document

    if "*" in line:
        words = line.split(" ")

        p = document.add_paragraph(style = style)

        for word in words:
            word += " "

            if word[0:2] == "**":
                word = word.replace("**", "")

                p.add_run(word).bold = True
            elif word[0] == "*":
                word = word.replace("*", "")

                p.add_run(word).italic = True
            else:
                p.add_run(word)
    else:
        document.add_paragraph(line, style = style)

# Doc Part

def add_part(filename):
    global document

    with open(f"docs/parts/{filename}", "r") as file:
        for line in file:
            line = line.strip()

            if line != "":
                first = line[0]

                if first == "#":
                    parts = line.split(" ")

                    level = 1

                    if line[0:2] == "##": level = 2

                    line = re.sub("#\s?", "", line)

                    document.add_heading(line, level)

                    document.add_paragraph()
                elif first == "-":
                    line = line.replace("- ", "")

                    add_content(line, style = "List Bullet")
                elif first == "%":
                    line = line.replace("%", "")

                    parts = line.split(";")

                    filename = parts[0]
                    num_cols = int(parts[1])

                    add_table(filename, num_cols)

                    document.add_paragraph()
                elif first == "*":
                    line = line.replace("*", "")

                    p = document.add_paragraph()
                    p.add_run(line).bold = True
                elif first == "!":
                    filename = line[1:].replace("\"", "")
                    
                    add_image(filename)
                else:
                    add_content(line)

        document.add_page_break()

        file.close()

# Docx file
document = Document()

# Set font
style = document.styles["Normal"]
font = style.font

font.name = "Calibri"

# Add parts
#files = ["EN-resume.md", "context-projet.md"]

#add_part("_test.md")
#add_part("EN-resume.md")
#add_part("expression-besoins.md")
add_part("base-donnees.md")

#add_part("spe-fonctionnelles.md")
#add_part("base-donnees.md")

# Compilation
document.save("compilation.docx")

print("Compilation du dossier effectuée")